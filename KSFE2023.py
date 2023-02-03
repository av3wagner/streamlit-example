#!/usr/bin/env python
# coding: utf-8

# https://stackoverflow.com/questions/32932230/add-an-image-in-a-specific-position-in-the-document-docx

# ### Start: 20.11.2022<br> 
# ### Version: 03.02.2023<br> <br>
# ### D:\IPYNB\KSFE2023\PROGRAMME\POWERSHELL_PYTHON.au3<br> 
# %windir%\System32\cmd.exe "/K" C:\ANACONDA\Scripts\activate.bat<br> 
# Python D:\IPYNB\KSFE2023\PROGRAMME\AddImageUndTables.py<br> 
# ### Header and Footer<br>
# https://docstest.readthedocs.io/en/latest/dev/analysis/features/header.html<br>
# 
# ### Autoit-Script<br> 
# While (FileExists("D:\IPYNB\KSFE2023\OUTPUT\AddImage.docx"))<br> 
# 	FileDelete("D:\IPYNB\KSFE2023\OUTPUT\AddImage.docx")<br> 
# 	Sleep(50)<br> 
# WEnd<br> 
# Send("#r")<br> 
# Sleep(50)<br> 
# Send('%windir%\System32\cmd.exe "/K" C:\ANACONDA\Scripts\activate.bat')<br> 
# Sleep(50)<br> 
# Send("{Enter}")<br> 
# Sleep(500)<br> 
# <br> 
# Send('Python D:\IPYNB\KSFE2023\PROGRAMME\AddImageUndTables.py')<br> 
# Sleep(50)<br> 
# Send("{Enter}")<br> 
# 
# While NOT (FileExists("D:\IPYNB\KSFE2023\OUTPUT\AddImage.docx"))<br> 
# 	Sleep(50)<br> 
# WEnd<br> 
# 
# OpenWord()<br> 
# ### Google: python-docx style rtl mod!!!<br>
# 
# ### Word-Open
# 
# https://stackoverflow.com/questions/52889704/python-win32com-excel-com-model-started-generating-errors<br>
# ### from pathlib import Path
# ### import win32com
# ### gen_py_path = 'C:/User/satur/AppData/Local/AW/gen_py'
# 
# ### Path(gen_py_path).mkdir(parents=True, exist_ok=True)
# ### win32com.__gen_path__ = gen_py_path
# ### print(win32com.__gen_path__)
# ### Reult: C:/User/satur/AppData/Local/AW/gen_py

# In[ ]:


import sys
import inspect, os
#import win32com.client
from win32com.client import constants
import win32com.client, time, pythoncom
from win32com import client
import requests
import base64

import re
import atexit
from time import time, strftime, localtime
from datetime import timedelta
from copy import deepcopy
import pandas as pd
from pathlib import Path


# In[ ]:


import docx
from docx import *
from docx import Document, enum
from docx.oxml import *
from docx.shared import Pt, Mm, Cm, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.enum.table import WD_TABLE_ALIGNMENT 
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_BUILTIN_STYLE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# In[ ]:


#os.chdir("C:/IPYNB/KSFE2023/")
githubToken = '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'
path = os.getcwd()
print ("Directory changed successfully %s" % path)
docx_file=os.path.join(path, "", "OUTPUT", "KSFE2023.docx")
docx_Out=os.path.join(path, "", "OUTPUT", "KSFE2023.pdf")
print("Doc Infile", docx_file)
print("Doc Outfile",docx_Out)

if os.path.isfile(docx_file):
    os.remove(docx_file)
else:
    print("Error: %s file not found" % docx_file)


# In[ ]:


form = Document(docx = ".\INDATEN\Tables.docx")
df = pd.read_csv('.\\DATEN\\Export.csv')
country = pd.read_csv('.\\INDATEN\\Country1.csv')
input_doc = Document('.\\INDATEN\\Para1.docx')
front_doc = Document('.\\INDATEN\\front.docx')
text_file = open('.\\INDATEN\\TxtGrVor.txt', "r")

df3=df[(df['GR']==2)]
df3=df3.drop(['GR','NN'], axis=1)
Kopf = ['red', '#9C0C38', 'orange']
n_rows, n_cols = df3.shape[0], df3.shape[1] 

data = text_file.read()
text_file.close()


# In[ ]:


def add_content(content, color, space_after, font_name='Arial', font_size=16, line_spacing=0, space_before=0,
                align='WD_ALIGN_PARAGRAPH.CENTER', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=True,style_name=""):
    paragraph = doc.add_paragraph(content)
    paragraph.style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = set_bold
    font.italic = set_italic
    font.all_caps = set_all_caps
    font.underline = set_underline
    font.color.rgb = color
    paragraph_format = paragraph.paragraph_format
    paragraph.alignment = align
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = line_spacing
    paragraph_format.keep_together = keep_together
    paragraph_format.keep_with_next = keep_with_next
    paragraph_format.page_break_before = page_break_before
    paragraph_format.widow_control = widow_control

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    return output_para


# In[ ]:


def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT #LEFT #CENTER #RIGHT
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    #t1.text = 'Erstellt von Dr. Alexander Wagner: Page '
    t1.text = 'Page '
    page_run._r.append(t1)
    page_num_run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')
    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    
    instrText2.text = "NUMPAGES"
    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')
    
    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

def tab_copy(k): 
    s=22
    p2 = doc.add_heading("Tabelle №" + str(s+k+1) , 3) 
    table = form.tables[k]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    paragraph = doc.add_paragraph("Hier wird die Tabelle №" + str(s+k+1) + " kopiert! Alle Daten wurde Simuliert!")
    paragraph._p.addnext(tbl)
    
def dtoc():
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    doc.add_section()
    
def set_table_header_bg_color(tc):
    """
    set background shading for Header Rows
    """
    tblCellProperties = tc._element.tcPr
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), "0099CC")   # 0,153,204
    tblCellProperties.append(clShading)    
    
def Delete_table(table):
        doc.tables[table]._element.getparent().remove(doc.tables[table]._element)     


# In[ ]:


def update_toc0(docx_file):
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = 1  
    word.DisplayAlerts = 0

    doc = word.Documents.Open(docx_file)
    wd_section = doc.Sections(1)  
    toc_count = doc.TablesOfContents.Count
    print(toc_count)
    stringG='INHALTSVERZEICHNIS'
    stringK='Inhaltsverzeichnis'
    if toc_count == 0:
        for i, p in enumerate(doc.Paragraphs):
            if stringK in p.Range.Text:
                try:
                    p.Range.InsertParagraphAfter()
                    parag_range = doc.Paragraphs(i+2).Range
                    parag_range.Font.Name = 'Arial'
                    parag_range.Font.Size = 14
                    parag_range.Font.Bold = constants.wdToggle
                    parag_range.Font.Size = 12
                    doc.TablesOfContents.Add(Range=parag_range,
                                             UseHeadingStyles=True,
                                             LowerHeadingLevel=3)
                except Exception as e:
                    print("Ja：", e, "Nein")
                break

    elif toc_count == 1:
        toc = doc.TablesOfContents(1)
        toc.Update()
        print('TOC should have been updated.')
    else:
        print('TOC has not been updated for sure...')      
    


# https://stackoverflow.com/questions/1557571/how-do-i-get-time-of-a-python-programs-execution

# In[ ]:


def secondsToStr(elapsed=None):
    if elapsed is None:
        #return strftime("%Y-%m-%d %H:%M:%S", localtime())
        return strftime("%d.%m.%Y %H:%M:%S", localtime())
    else:
        return str(timedelta(seconds=elapsed))

def log(s, elapsed=None):
    line = "="*40
    print(s,  secondsToStr())
    if elapsed:
        print("Elapsed time:", elapsed)
    print()

def endlog():
    end = time()
    elapsed = end-start
    log("End Program", secondsToStr(elapsed))


# In[ ]:


def xpage():
    doc = Document()  
    styles = doc.styles
    rtlstyle = doc.styles.add_style('rtl', enum.style.WD_STYLE_TYPE.PARAGRAPH)

    print(rtlstyle.name)
    print(rtlstyle.font.name)
    print(rtlstyle.font.size)
    section = doc.sections[0] 
    print('Section layout start', section.page_width.cm, section.page_height.cm)
    print("Orientation start:", section.orientation)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    font.size = docx.shared.Pt(11)

    print(style.name)
    print(style.font.name)
    print(style.font.size)

    new_string = 'Inhaltsverzeichnis'    
    p=doc.add_paragraph(new_string)  
    section = doc.add_section()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)
    print("Ende def page!")


# In[ ]:


def set_repeat_table_header(row):
   """ set repeat table row on every new page
   """
   tr = row._tr
   trPr = tr.get_or_add_trPr()
   tblHeader = OxmlElement('w:tblHeader')
   tblHeader.set(qn('w:val'), "true")
   trPr.append(tblHeader)
   return row


# In[ ]:


def change_table_cell(cell, background_color=None, font_color=None, font_size=None, bold=None, italic=None):
    """ changes the background_color or font_color or font style (bold, italic) of this cell.
    Leave the params as 'None' if you do not want to change them.
    params:
        cell: the cell to manipulate
        background_color: name for the color, e.g. "red" or "ff0000"
        font_color:
        font_size: size in pt (e.g. 10)
        bold:   requested font style. True or False, or None if it shall remain unchanged
        italic: requested font style. True or False, or None if it shall remain unchanged
    background_color: the color of cells background"""
    if background_color:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    if font_color:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = docx.shared.RGBColor.from_string(font_color)

    if font_size:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = docx.shared.Pt(font_size)

    if bold is not None:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = bold

    if italic is not None:
        for p in cell.paragraphs:
            for r in p.runs:
                r.italic = italic

def change_table_row(table_row, background_color=None, font_color=None, font_size=None, bold=None, italic=None):
    for cell in table_row.cells:
        change_table_cell(cell, background_color=background_color, font_color=font_color, font_size=font_size,
                          bold=bold,
                          italic=italic)


# In[ ]:


def xtables():
    widths = [Cm(6), Cm(1.8), Cm(1.7), Cm(2), Cm(2), Cm(2), Cm(2)]
    include_index=True 
    for k in range (0, 22): 
        df3=df[(df['GR']==k+1)]
        df3=df3.drop(['GR','NN'], axis=1)
        n_rows, n_cols = df3.shape[0], df3.shape[1]  
        picPath = ".\\OUTPUT\\GRAPH" + str(k+1) + ".png"
        GraphKopf="Grafik №"+str(k+1)

        doc.add_heading("Bericht Block №" +str(k+1)) 
        doc.add_heading("Text Block №"+str(k+1),  2) 
        p = doc.add_paragraph("Hinweis: Alle Text Blöcke wurde vom Python-Dokumentation kopiert!")
        for run in p.runs:
            run.bold = True
            run.underline = True
            run.italic = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            
        for j in range(0, 4):
            get_para_data(doc, input_doc.paragraphs[k*15+j])
            
        doc.add_heading("Tabelle №"+str(k+1), 2) 
        doc.add_heading("Altersverteilung für ausgewählte Länder nach WHO: " + str(country.values[k, 1]), 3) 
        t = doc.add_table(n_rows+2, n_cols, style="Table Grid")
        #t.allow_autofit = True
        set_repeat_table_header(t.rows[0])
        for j in range(n_cols): 
            if include_index:
                ha=1
            else:
                t.cell(0,j).text = df3.columns[j]
                
        for i in range(2, n_rows+2): 
            for j in range(n_cols): 
                if include_index:
                    t.cell(i, j).text = str(df3.values[i-2,j])
                    t.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
                else:
                    t.cell(i, j).text = str(df3.values[i,j])
            t.cell(i, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT 
            t.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
        t.allow_autofit = True    
        t.style = 'Colorful Shading Accent 1' 

        column_names0 = ["Country", "Median"]
        column_names = ["0-14", "15-24", "25-54", "55-64", "65+"]
        
        heading_cells = t.rows[0].cells
        for j in range(0, 1):
            heading_cells[j].text = column_names0[j]
            heading_cells[j].width = widths[j]
        
        heading_cells = t.rows[1].cells
        for j in range(0, len(column_names)):
            heading_cells[j+2].text = column_names[j]
            heading_cells[j].width = widths[j]
            cell = t.cell(1, j+2)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
            
        for j in range(0, n_cols):
            set_table_header_bg_color(t.rows[0].cells[j])
        
        a = t.cell(0, 0)
        a.text = " "  
        b = t.cell(1, 0)
        b.merge(a)
        
        cell = t.cell(0, 0)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        #run.add_picture("D:\\IPYNB\\KSFE2023\\INDATEN\\who-logo-png.png", Cm(1), Cm(1))
        run.add_picture(".\\INDATEN\\who-logo-png.png", Cm(1), Cm(1))
        
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER   
          
        cell = t.cell(1, 1)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
        cell.paragraphs[0].add_run("Median")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  
     
        cell = t.cell(0, 1)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
        cell.paragraphs[0].add_run("Gesamt Population")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  
        change_table_cell(t.rows[0].cells[1], font_color="ffc000", font_size=8, bold=True, italic=True)
         
        e = t.cell(0, 2)
        f = t.cell(0, 6)
        e.merge(f)
        
        cell = t.cell(0, 6)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        cell.paragraphs[0].add_run("Altersgruppen")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER 
           
        t.add_row()
        g = t.cell(n_rows+2, 1)
        h = t.cell(n_rows+2, 6)
        g.merge(h)
        cell = t.cell(n_rows+2, 6)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 
        cell.paragraphs[0].add_run("© Dr. Alexander Wagner. All Rights Reserved")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER 
        change_table_cell(t.rows[n_rows+2].cells[2], background_color="lightgreen", font_color="0000ff", font_size=8, bold=True, italic=True)
                          
        p = doc.add_paragraph(" ")
        p = doc.add_paragraph("Hinweis: Alle Text Blöcke wurde vom Python-Dokumentation kopiert!")
        for run in p.runs:
            run.bold = True
            run.underline = True
            run.italic = True
            run.font.color.rgb = RGBColor(0x7A, 0x04, 0x92)
            
        for j in range(5, 9):
            get_para_data(doc, input_doc.paragraphs[k*15+j])    
            
        doc.add_heading(GraphKopf, 2) 
        doc.add_heading("Altersverteilung für ausgewählte Länder nach WHO: " + str(country.values[k, 1]), 3) 
        
        p = doc.add_paragraph(" ")
        p = p.insert_paragraph_before('')
        r = p.add_run()
        r.add_picture(picPath, width=Cm(18))
        
        p = doc.add_paragraph("Hinweis: Alle Text Blöcke wurde vom Python-Dokumentation kopiert!")
        for run in p.runs:
            run.bold = True
            run.underline = True
            run.italic = True
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            #run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        for j in range(10, 14):
            get_para_data(doc, input_doc.paragraphs[k*15+j])
   
        doc.add_section()
    doc.add_heading("Bericht Block №23", 1) 
      
    for tb in range(8):  
        k=345
        s=22
        print(tb)
        doc.add_heading("Text Block №" + str(s+1+tb), 3) 
        p = doc.add_paragraph("Hinweis: Alle Text Blöcke wurde vom Python-Dokumentation kopiert!")
        for run in p.runs:
            run.bold = True
            run.underline = True
            run.italic = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        
        for j in range(0, 4):
            get_para_data(doc, input_doc.paragraphs[k+tb*5+j])
            #tab_copy(tb)
        if tb < 7:
            p2 = doc.add_paragraph("Die Tabelle №"  + str(s+tb+1) + " wurde erfolgreich kopiert!")
            p2 = doc.add_paragraph()
            doc.add_section()
     
    for nt in range(29, 32):
        print("Sections №", nt)    
        section = doc.sections[nt]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_height = Mm(210)
        section.page_width = Mm(297)    
  
    #print("Ende def tables!")


# https://github.com/python-openxml/python-docx/issues/957<br>

# In[ ]:


def post_page():
    doc = Document( ) 
    header = section.header
    w=section.page_width.cm-4
    hetable = header.add_table(1, 2, Cm(w))
    
    cell = hetable.rows[0].cells[1]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(".\\INDATEN\\KI3.jpg")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cell = hetable.rows[0].cells[0]
    paragraph = cell.paragraphs[0]    
    paragraph.text = "Projektinitiative: Ein Web- und Cloudbasiertes Multiple-Kernel Eco-System für die automatisierte Erstellung von analytischen Berichten.\nEntwickelt von Dr. Alexander Wagner im November 2022"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.style = rtlstyle 
    font = paragraph.style.font
    font.rtl = False
    font.name = 'Times New'
    font.size = Pt(8)
    font.bold = True
    font.italic = True

    #Footer
    footer = section.footer
    for paragraph in footer.paragraphs:
        paragraph.clear

    log("Programm endet:")
    version=3
    dato=secondsToStr() 
    docOwner='Dr. Alexander Wagner'

    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font = footer_para.style.font
    font.name = 'Times New'
    font.size = Pt(8)
    font.bold = True
    font.italic = True

    footer_para.add_run('Datum der Berichtserstellung: ' + dato + ', Programm-Version: №' + str(version) + '\t\t') 
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.add_run('\nDocument owner: ' + docOwner + '\t\t')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.add_run("\n© Paradox GmbH. All Rights Reserved \t\t\t\t")
    add_page_number(footer_para)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    print(section.different_first_page_header_footer)
    section.different_first_page_header_footer = True
    print(section.different_first_page_header_footer)            


# In[ ]:


def open_doc():
    update_toc(docx_file)
    print("Ende def open_doc!")


# In[ ]:


doc = Document( ) 
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(20)
section.right_margin = Mm(20)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.header_distance = Mm(7)
section.footer_distance = Mm(7)

txt=""
align='WD_ALIGN_PARAGRAPH.CENTER'
font_styles = doc.styles
font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(20)
font_object.color.rgb = RGBColor(0x2F, 0x2F, 0xFF)
font_object.name = 'Times New Roman'

parag = doc.add_paragraph(txt)
parag.add_run('Projektinitiative', style='CommentsStyle').bold = True
parag.alignment = WD_ALIGN_PARAGRAPH.CENTER
parag = doc.add_paragraph(txt)

font_styles = doc.styles
font_charstyle = font_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(24)
font_object.color.rgb = RGBColor(0x00, 0x00, 0x00)
font_object.name = 'Times New Roman'
parag.add_run('Ein Web- und Cloudbasiertes Multiple-Kernel \n Eco-System für die automatisierte Erstellung \n von analytischen Berichten', style='CommentsStyle1').bold = True
parag.alignment = WD_ALIGN_PARAGRAPH.CENTER

for t in range(0, 4):
    p=doc.add_paragraph(txt) 

picPath = ".\\INDATEN\\KI.jpg"    
p = doc.add_paragraph(" ")
r = p.add_run()
r.add_picture(picPath, Cm(10), Cm(6))
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('')
r = p.add_run() 

for t in range(0,8):
    txt=""
    p=doc.add_paragraph(txt) 
    
txt="Berlin"
add_content(txt, color=RGBColor.from_string('0000FF'), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=20, line_spacing=1, font_name='Time New', font_size=14,
            set_bold=True, set_all_caps=True,style_name ="NormalCENTERS ")

txt="2023"
add_content(txt, color=RGBColor.from_string('0000FF'), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10, line_spacing=1, font_name='Time New', font_size=12,
            set_bold=True, set_all_caps=True,style_name ="NormalCENTER")


section = doc.add_section()
new_string = 'Inhaltsverzeichnis'    
p=doc.add_paragraph(new_string)  
for run in p.runs:
    run.font.name = 'Times New'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
    
dtoc()
section = doc.add_section()

styles = doc.styles
rtlstyle = doc.styles.add_style('rtl', enum.style.WD_STYLE_TYPE.PARAGRAPH)
rtlstyle.font.rtl = True
section = doc.sections[0] 
print('Section layout start', section.page_width.cm, section.page_height.cm)
print("Orientation start:", section.orientation)

font_styles = doc.styles
font_charstyle = font_styles.add_style('Cambria', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(11)
font_object.name = 'Times New Roman'


# In[ ]:


def update_toc(docx_file):
    word = client.DispatchEx("Word.Application") 
    doc = word.Documents.Open(docx_file)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()
    
def xmain(docx_file):
    update_toc(docx_file)
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = 1  
    word.DisplayAlerts = 0
    doc = word.Documents.Open(docx_file)


# In[ ]:


xtables()
post_page() 
doc.save(docx_file)
print("Ende def post_page!")


# In[ ]:


wdFormatPDF = 17
xlTypePDF = 0
ppSaveAspdf = 32

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
word.DisplayAlerts = False
doc = word.Documents.Open(docx_file)
doc.SaveAs(docx_Out, wdFormatPDF)
doc.Close()
word.Quit()    


# https://docs.github.com/de/rest/repos/contents?apiVersion=2022-11-28#create-or-update-file-contents<br>
# https://api.github.com/repos/av3wagner/KSFE2023/contents/KSFE2023.docx

# In[ ]:


def push(file):
    with open(file, "rb") as f:        
        encodedData = base64.b64encode(f.read())
        headers = {
            "Authorization": f'''Bearer {githubToken}''',
            "Content-type": "application/vnd.github+json"
        }
        data = {
            "message": "My commit message",             # Put your commit message here.
            "content": encodedData.decode("utf-8")      # Super: encodedData.decode !!!
        }
        r = requests.put(githubAPIURL, headers=headers, json=data)
        print(r.text) 

udir = "https://api.github.com/repos/av3wagner/KSFE2023/contents/"
print (udir)
ulist = ["KSFE2023.docx", "KSFE2023.pdf"]

for l in range(len(ulist)):
    githubAPIURL = udir + ulist[l]
    print(githubAPIURL)
    file=os.path.join(os.getcwd(), "", "OUTPUT", ulist[l])                           
    print(file)
    push(file)


# In[ ]:




