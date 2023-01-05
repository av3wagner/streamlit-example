import sys
import inspect, os
import io
from collections import namedtuple
import altair as alt
import math
import pandas as pd
import streamlit as st
import subprocess
import numpy as np

df = pd.DataFrame(np.random.randn(500, 2) / [50, 50] + [37.76, -122.4],
columns=['lat', 'lon'])
st.map(df)
current_dir = ".\" 
#r"D:\IPYNB\KSFE2023\PROGRAMME"
subprocess.Popen(os.path.join(current_dir,"RUN_VAS2023_TestD.exe"))
#subprocess.Popen(".\RUN_VAS2023_TestD.exe")

'''
#doc_file=".\data\StApp.docx"
#doc_download = doc_file_creation(doc_file)
doc_download=1
bio = io.BytesIO()
#doc_download.save(bio)
if doc_download:
    st.download_button(
    label="Click here to download",
    data=bio.getvalue(),
    file_name="StApp.docx",
     mime="docx"
     )

import docx
from docx import *
from docx import Document, enum
from copy import deepcopy
import pandas as pd
from pathlib import Path

from docx.oxml import *
from docx.shared import Pt, Mm, Cm, Inches, RGBColor
#from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT

from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.enum.table import WD_TABLE_ALIGNMENT 
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
#from docx.shared import Pt
from docx.enum.style import WD_BUILTIN_STYLE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import sys
import inspect, os
import win32com.client
from win32com.client import constants
import re
import atexit
from time import time, strftime, localtime
from datetime import timedelta


# In[63]:


def secondsToStr(elapsed=None):
    if elapsed is None:
        #return strftime("%Y-%m-%d %H:%M:%S", localtime())
        return strftime("%d.%m.%Y %H:%M:%S", localtime())
    else:
        return str(timedelta(seconds=elapsed))

def log(s, elapsed=None):
    line = "="*40
    print(s,  secondsToStr())
    if elapsed:
        print("Elapsed time:", elapsed)
    print()

def endlog():
    end = time()
    elapsed = end-start
    log("End Program", secondsToStr(elapsed))


# In[64]:


def add_content(content, color, space_after, font_name='Arial', font_size=16, line_spacing=0, space_before=0,
                align='WD_ALIGN_PARAGRAPH.CENTER', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=True,style_name=""):
    paragraph = doc.add_paragraph(content)
    paragraph.style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = set_bold
    font.italic = set_italic
    font.all_caps = set_all_caps
    font.underline = set_underline
    font.color.rgb = color
    paragraph_format = paragraph.paragraph_format
    #paragraph_format.alignment = alignment_dict.get(align.lower())
    paragraph.alignment = align
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = line_spacing
    paragraph_format.keep_together = keep_together
    paragraph_format.keep_with_next = keep_with_next
    paragraph_format.page_break_before = page_break_before
    paragraph_format.widow_control = widow_control

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    #element.set(nsqn(name), value)
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT #LEFT #CENTER #RIGHT
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)
    page_num_run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')
    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    
    instrText2.text = "NUMPAGES"
    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')
    
    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)


def post_page():
    doc = Document( ) 
    #header
    header = section.header
    w=section.page_width.cm-4
    #print(w)
    hetable = header.add_table(1, 2, Cm(w))
    
    cell = hetable.rows[0].cells[1]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(".\data\KI3.jpg")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cell = hetable.rows[0].cells[0]
    paragraph = cell.paragraphs[0]    
    paragraph.text = "Projektinitiative: Ein Web- und Cloudbasiertes Multiple-Kernel Eco-System für die automatisierte Erstellung von analytischen Berichten.\nEntwickelt von Dr. Alexander Wagner im Januar 2023"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.style = rtlstyle 
    font = paragraph.style.font
    font.rtl = False
    font.name = 'Times New'
    font.size = Pt(8)
    font.bold = True
    font.italic = True

    #Footer
    footer = section.footer
    for paragraph in footer.paragraphs:
        paragraph.clear

    log("Programm endet:")
    version=3
    dato=secondsToStr() 
    docOwner='Dr. Alexander Wagner'

    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font = footer_para.style.font
    font.name = 'Times New'
    font.size = Pt(8)
    font.bold = True
    font.italic = True

    footer_para.add_run('Datum der Berichtserstellung: ' + dato + ', Programm-Version: №' + str(version) + '\t\t') 
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.add_run('\nDocument owner: ' + docOwner + '\t\t')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.add_run("\n© Paradox GmbH. All Rights Reserved \t\t\t\t")
    add_page_number(footer_para)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    print(section.different_first_page_header_footer)
    section.different_first_page_header_footer = True
    print(section.different_first_page_header_footer)  
          

doc = Document( ) 
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(20)
section.right_margin = Mm(20)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.header_distance = Mm(7)
section.footer_distance = Mm(7)

txt=""
align='WD_ALIGN_PARAGRAPH.CENTER'
font_styles = doc.styles
font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(20)
#font_object.color.rgb = RGBColor(0x00, 0xFF, 0x00)
font_object.color.rgb = RGBColor(0x2F, 0x2F, 0xFF)
font_object.name = 'Times New Roman'


parag = doc.add_paragraph(txt)
parag.add_run('Projektinitiative', style='CommentsStyle').bold = True
parag.alignment = WD_ALIGN_PARAGRAPH.CENTER
parag = doc.add_paragraph(txt)

font_styles = doc.styles
font_charstyle = font_styles.add_style('CommentsStyle1', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(24)
font_object.color.rgb = RGBColor(0x00, 0x00, 0x00)
font_object.name = 'Times New Roman'
parag.add_run('Ein Web- und Cloudbasiertes Multiple-Kernel \n Eco-System für die automatisierte Erstellung \n von analytischen Berichten', style='CommentsStyle1').bold = True
parag.alignment = WD_ALIGN_PARAGRAPH.CENTER

for t in range(0, 4):
    p=doc.add_paragraph(txt) 

picPath = ".\data\KI.jpg"    
p = doc.add_paragraph(" ")
r = p.add_run()
r.add_picture(picPath, Cm(10), Cm(6))
r.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('')
r = p.add_run() 

for t in range(0,8):
    txt=""
    p=doc.add_paragraph(txt) 
    
txt="Berlin"
add_content(txt, color=RGBColor.from_string('0000FF'), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=20, line_spacing=1, font_name='Time New', font_size=14,
            set_bold=True, set_all_caps=True,style_name ="NormalCENTERS ")

txt="2023"
add_content(txt, color=RGBColor.from_string('0000FF'), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=10, line_spacing=1, font_name='Time New', font_size=12,
            set_bold=True, set_all_caps=True,style_name ="NormalCENTER")

section = doc.add_section()
new_string = 'Inhaltsverzeichnis'    
p=doc.add_paragraph(new_string)  
for run in p.runs:
    run.font.name = 'Times New'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
    

section = doc.add_section()

styles = doc.styles
rtlstyle = doc.styles.add_style('rtl', enum.style.WD_STYLE_TYPE.PARAGRAPH)
rtlstyle.font.rtl = True
print(rtlstyle.name)
print(rtlstyle.font.name)
print(rtlstyle.font.size)
section = doc.sections[0] 
print('Section layout start', section.page_width.cm, section.page_height.cm)
print("Orientation start:", section.orientation)

font_styles = doc.styles
font_charstyle = font_styles.add_style('Cambria', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(11)
font_object.name = 'Times New Roman'

post_page()
docx_file=".\data\StApp.docx"
doc.save(docx_file)
os.startfile(".\data\StApp.docx")
'''





